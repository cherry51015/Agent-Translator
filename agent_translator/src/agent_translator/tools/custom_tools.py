"""
Custom Tools for CrowdWisdomTrading Document Translation System
Standalone version without crewai_tools dependency
"""

import os
import json
import re
import time
from typing import Dict
from pathlib import Path

from docx import Document
import yaml
import litellm
from litellm import completion


# Minimal replacement for missing BaseTool/Tool
class SimpleTool:
    """Minimal replacement for CrewAI Tool/BaseTool"""
    name: str = "Unnamed Tool"
    description: str = "No description"

    def run(self, *args, **kwargs):
        return self._run(*args, **kwargs)

    def _run(self, *args, **kwargs):
        raise NotImplementedError("Tool must implement _run()")


class DocumentAnalyzerTool(SimpleTool):
    name = "Document Analyzer Tool"
    description = "Analyzes document type, academic field, language, and structure"

    def _run(self, document_path: str) -> str:
        """Analyze document and return comprehensive analysis"""
        try:
            if not os.path.exists(document_path):
                raise FileNotFoundError(f"Document not found: {document_path}")

            doc = Document(document_path)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            document_content = "\n".join(full_text)
            word_count = len(document_content.split())

            analysis_prompt = f"""
            Analyze this academic document excerpt (first 1000 characters):

            {document_content[:1000]}

            Determine:
            1. Document type (research paper, review, case study, etc.)
            2. Academic field (medicine, engineering, computer science, etc.)
            3. Source language
            4. Complexity level
            5. Special elements present

            Respond in exact JSON format:
            {{
                "document_type": "type",
                "academic_field": "field",
                "source_language": "language",
                "complexity_level": "level",
                "special_elements": ["element1", "element2"],
                "translation_instructions": "field-specific translation guidance",
                "estimated_length": "{word_count} words"
            }}
            """

            response = completion(
                model="gemini/gemini-1.5-flash",
                messages=[{"role": "user", "content": analysis_prompt}],
                temperature=0.3,
            )
            return response.choices[0].message.content

        except Exception as e:
            return json.dumps({
                "error": f"Analysis failed: {str(e)}",
                "document_type": "unknown",
                "academic_field": "general",
                "source_language": "unknown",
                "complexity_level": "moderate",
                "special_elements": [],
                "translation_instructions": "Standard academic translation",
                "estimated_length": "unknown"
            })


class DocumentStructureTool(SimpleTool):
    name = "Document Structure Tool"
    description = "Extracts document structure and creates translation batches"

    def _run(self, document_path: str, max_tokens_per_batch: int = 2000) -> str:
        try:
            doc = Document(document_path)
            structure_data = {
                "total_paragraphs": 0,
                "formatting_map": {},
                "translation_batches": []
            }

            paragraphs_data = []
            current_batch = []
            current_tokens = 0
            batch_id = 1

            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue

                paragraph_data = {
                    "id": i,
                    "text": paragraph.text,
                    "formatting": self._extract_formatting(paragraph)
                }
                paragraphs_data.append(paragraph_data)
                structure_data["formatting_map"][str(i)] = paragraph_data["formatting"]

                # Estimate tokens
                paragraph_tokens = len(paragraph.text.split()) * 0.75

                if current_tokens + paragraph_tokens > max_tokens_per_batch and current_batch:
                    structure_data["translation_batches"].append({
                        "batch_id": f"batch_{batch_id}",
                        "paragraphs": current_batch.copy(),
                        "token_estimate": int(current_tokens),
                        "context_note": f"Batch {batch_id} of academic content"
                    })
                    current_batch = []
                    current_tokens = 0
                    batch_id += 1

                current_batch.append({
                    "paragraph_id": i,
                    "text": paragraph.text,
                    "formatting": paragraph_data["formatting"]
                })
                current_tokens += paragraph_tokens

            if current_batch:
                structure_data["translation_batches"].append({
                    "batch_id": f"batch_{batch_id}",
                    "paragraphs": current_batch,
                    "token_estimate": int(current_tokens),
                    "context_note": f"Final batch {batch_id} of academic content"
                })

            structure_data["total_paragraphs"] = len(paragraphs_data)
            return json.dumps(structure_data, indent=2)

        except Exception as e:
            return json.dumps({"error": f"Structure analysis failed: {str(e)}"})

    def _extract_formatting(self, paragraph) -> Dict:
        formatting = {"bold": [], "italic": [], "underline": []}
        for run in paragraph.runs:
            if run.bold:
                formatting["bold"].append(run.text)
            if run.italic:
                formatting["italic"].append(run.text)
            if run.underline:
                formatting["underline"].append(run.text)
        return formatting


class BatchTranslatorTool(SimpleTool):
    name = "Batch Translator Tool"
    description = "Translates content in intelligent batches with academic quality"

    def _run(self, structural_analysis: str, target_language: str, translation_instructions: str = "") -> str:
        try:
            structure_data = json.loads(structural_analysis)
            translation_batches = structure_data.get("translation_batches", [])
            if not translation_batches:
                raise ValueError("No translation batches found")

            translated_batches = []

            for batch in translation_batches:
                batch_content = [f"[PARA_{p['paragraph_id']}] {p['text']}" for p in batch["paragraphs"]]
                batch_text = "\n\n".join(batch_content)

                translation_prompt = f"""
                You are a professional academic translator. Translate the following content to {target_language}.

                REQUIREMENTS:
                1. Maintain academic tone and precision
                2. Preserve technical terminology appropriately
                3. Keep paragraph markers [PARA_X] intact
                4. Ensure academic concepts are accurately conveyed

                {translation_instructions}

                Content to translate:
                {batch_text}

                Provide ONLY the translated text with paragraph markers preserved.
                """

                response = completion(
                    model="gemini/gemini-1.5-flash",
                    messages=[{"role": "user", "content": translation_prompt}],
                    temperature=0.3,
                    max_tokens=4000
                )
                translated_content = response.choices[0].message.content

                translated_paragraphs = []
                for line in translated_content.split('\n\n'):
                    if line.strip() and '[PARA_' in line:
                        para_match = re.match(r'\[PARA_(\d+)\] (.*)', line.strip(), re.DOTALL)
                        if para_match:
                            para_id = int(para_match.group(1))
                            para_text = para_match.group(2)
                            translated_paragraphs.append({
                                "paragraph_id": para_id,
                                "original": next(p["text"] for p in batch["paragraphs"] if p["paragraph_id"] == para_id),
                                "translated": para_text,
                                "formatting": next(p["formatting"] for p in batch["paragraphs"] if p["paragraph_id"] == para_id)
                            })

                translated_batches.append({
                    "batch_id": batch["batch_id"],
                    "paragraphs": translated_paragraphs,
                    "status": "completed"
                })
                time.sleep(1)

            return json.dumps({
                "status": "completed",
                "total_batches": len(translated_batches),
                "translated_batches": translated_batches
            }, indent=2)

        except Exception as e:
            return json.dumps({"error": f"Translation failed: {str(e)}", "status": "failed"})


class DocxFormatterTool(SimpleTool):
    name = "DOCX Formatter Tool"
    description = "Creates formatted DOCX file from translated content"

    def _run(self, translated_content: str, structural_analysis: str, output_path: str) -> str:
        try:
            translation_data = json.loads(translated_content)
            if translation_data.get("status") != "completed":
                raise ValueError("Translation not completed successfully")

            doc = Document()
            all_paragraphs = {}
            for batch in translation_data["translated_batches"]:
                for para in batch["paragraphs"]:
                    all_paragraphs[para["paragraph_id"]] = para

            for para_id in sorted(all_paragraphs.keys()):
                translated_text = all_paragraphs[para_id]["translated"]
                paragraph = doc.add_paragraph()
                paragraph.add_run(translated_text)

            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)

            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path) / (1024 * 1024)
                return json.dumps({
                    "status": "completed",
                    "output_file": output_path,
                    "elements_recreated": {"paragraphs": len(all_paragraphs), "tables": 0, "footnotes": 0},
                    "file_size": f"{file_size:.2f} MB",
                    "quality_verification": "Document created successfully"
                }, indent=2)
            else:
                raise Exception("Failed to create output file")

        except Exception as e:
            return json.dumps({"status": "failed", "error": f"Document creation failed: {str(e)}", "output_file": None})
