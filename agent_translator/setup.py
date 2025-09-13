#!/usr/bin/env python3
"""
CrowdWisdomTrading Document Translation System Setup Script
This script will create the proper project structure and install dependencies
"""

import os
import sys
from pathlib import Path
import subprocess
import yaml

def create_project_structure():
    """Create the required project directory structure"""
    print("Creating project structure...")
    
    directories = [
        "src/agent_translator",
        "src/agent_translator/tools", 
        "src/agent_translator/config",
        "output",
        "logs"
    ]
    
    for directory in directories:
        Path(directory).mkdir(parents=True, exist_ok=True)
        print(f"  ✅ Created directory: {directory}")
    
    # Create __init__.py files
    init_files = [
        "src/agent_translator/__init__.py",
        "src/agent_translator/tools/__init__.py"
    ]
    
    for init_file in init_files:
        if not Path(init_file).exists():
            Path(init_file).touch()
            print(f"  ✅ Created file: {init_file}")

def create_config_files():
    """Create default configuration files"""
    print("Creating configuration files...")
    
    # Create config.yaml
    if not Path("config.yaml").exists():
        config_data = {
            "source_language": "English",
            "target_language": "Spanish", 
            "llm_model": "gemini/gemini-1.5-flash",
            "temperature": 0.3,
            "max_tokens_per_batch": 2000,
            "preserve_formatting": True,
            "output_directory": "output",
            "batch_processing": True,
            "translation_settings": {
                "maintain_academic_tone": True,
                "preserve_citations": True,
                "handle_technical_terms": True,
                "keep_formatting_markers": True
            }
        }
        
        with open("config.yaml", "w") as f:
            yaml.dump(config_data, f, default_flow_style=False, indent=2)
        print("  ✅ Created config.yaml")
    
    # Create .env template
    if not Path(".env").exists():
        env_content = """# CrowdWisdomTrading Translation System Environment Variables

# Required: Set one of these API keys
GEMINI_API_KEY=your_gemini_api_key_here
GOOGLE_API_KEY=your_google_api_key_here

# Optional: Alternative LLM providers
# OPENAI_API_KEY=your_openai_api_key_here
# ANTHROPIC_API_KEY=your_anthropic_api_key_here

# System Configuration
LOG_LEVEL=INFO
OUTPUT_DIR=output
"""
        with open(".env", "w") as f:
            f.write(env_content)
        print("  ✅ Created .env template")

def install_dependencies():
    """Install required Python packages"""
    print("Installing dependencies...")
    
    try:
        subprocess.check_call([
            sys.executable, "-m", "pip", "install", 
            "crewai[tools]==0.177.0",
            "litellm>=1.0.0", 
            "python-docx>=0.8.11",
            "PyYAML>=6.0",
            "click>=8.0.0",
            "structlog>=23.0.0"
        ])
        print("  ✅ Dependencies installed successfully")
        return True
    except subprocess.CalledProcessError as e:
        print(f"  ❌ Failed to install dependencies: {e}")
        return False

def write_init_files():
    """Write proper __init__.py files"""
    print("Writing __init__.py files...")
    
    # Main package init
    main_init_content = '''"""
CrowdWisdomTrading Document Translation System
AI-powered academic document translation with formatting preservation
"""

__version__ = "2.0.0"
__author__ = "CrowdWisdomTrading"

try:
    from .crew import DocumentTranslationCrew
    __all__ = ["DocumentTranslationCrew"]
except ImportError:
    # Handle case where crew.py doesn't exist yet
    __all__ = []
'''
    
    with open("src/agent_translator/__init__.py", "w") as f:
        f.write(main_init_content)
    
    # Tools package init
    tools_init_content = '''"""
Custom tools for document translation system
"""

try:
    from .custom_tools import (
        DocumentAnalyzerTool,
        DocumentStructureTool,
        BatchTranslatorTool,
        DocxFormatterTool
    )
    
    __all__ = [
        "DocumentAnalyzerTool",
        "DocumentStructureTool", 
        "BatchTranslatorTool",
        "DocxFormatterTool",
    ]
except ImportError:
    # Handle case where custom_tools.py doesn't exist yet
    __all__ = []
'''
    
    with open("src/agent_translator/tools/__init__.py", "w") as f:
        f.write(tools_init_content)
    
    print("  ✅ __init__.py files written")

def check_api_keys():
    """Check if API keys are configured"""
    print("Checking API keys...")
    
    api_keys = {
        "GEMINI_API_KEY": os.getenv("GEMINI_API_KEY"),
        "GOOGLE_API_KEY": os.getenv("GOOGLE_API_KEY"),
        "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY"),
        "ANTHROPIC_API_KEY": os.getenv("ANTHROPIC_API_KEY")
    }
    
    configured_keys = {k: v for k, v in api_keys.items() if v}
    
    if configured_keys:
        for key in configured_keys:
            print(f"  ✅ {key} is configured")
        return True
    else:
        print("  ❌ No API keys found")
        print("  📝 Please edit .env file and add your API key")
        return False

def create_sample_document():
    """Create a sample document for testing"""
    print("Creating sample document...")
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Sample Academic Document', 0)
        
        doc.add_paragraph(
            'This is a sample academic document for testing the '
            'CrowdWisdomTrading AI Translation System.'
        )
        
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            'Academic documents require careful translation to maintain '
            'their scholarly tone and technical precision. This system '
            'uses advanced AI to preserve formatting while ensuring '
            'accurate translation of complex concepts.'
        )
        
        doc.add_heading('Methodology', level=1) 
        doc.add_paragraph(
            'The translation process involves document analysis, structure '
            'breakdown, batch translation, and document reconstruction. '
            'Each step is handled by specialized AI agents working in '
            'coordination.'
        )
        
        doc.save('sample_document.docx')
        print("  ✅ Created sample_document.docx for testing")
        return True
        
    except ImportError:
        print("  ⚠️  Could not create sample document (python-docx not installed yet)")
        return False

def main():
    """Main setup function"""
    print("🚀 CrowdWisdomTrading Document Translation System Setup")
    print("=" * 60)
    
    # Step 1: Create project structure
    create_project_structure()
    print()
    
    # Step 2: Create configuration files
    create_config_files()
    print()
    
    # Step 3: Write __init__.py files
    write_init_files()
    print()
    
    # Step 4: Install dependencies
    if install_dependencies():
        print()
        
        # Step 5: Create sample document
        create_sample_document()
        print()
    
    # Step 6: Check API keys
    has_api_key = check_api_keys()
    print()
    
    # Final instructions
    print("📋 Setup Summary:")
    print("=" * 60)
    print("✅ Project structure created")
    print("✅ Configuration files created")
    print("✅ Dependencies installation attempted")
    
    if has_api_key:
        print("✅ API keys configured")
    else:
        print("❌ API keys need configuration")
    
    print("\n📝 Next Steps:")
    print("1. Edit .env file and add your API key")
    print("2. Copy the crew.py file to src/agent_translator/")
    print("3. Copy the custom_tools.py file to src/agent_translator/tools/")
    print("4. Copy the main.py file to the root directory")
    print("5. Run: python main.py test")
    print("6. Run: python main.py translate sample_document.docx")
    
    print("\n🎉 Setup completed! Follow the next steps to start translating.")

if __name__ == "__main__":
    main()