#!/usr/bin/env python3
"""
CrowdWisdomTrading Direct Translation System
Bypasses CrewAI complexity and does direct translation
"""

import os
import json
import logging
import argparse
import sys
import time
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import litellm
from litellm import completion

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('direct_translation.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class DirectDocumentTranslator:
    def __init__(self, target_language="Spanish"):
        self.target_language = target_language
        self.source_language = "English"
        self.model = "gemini/gemini-1.5-flash"
        self.temperature = 0.3
        self.total_tokens = 0
        
    def extract_document_content(self, file_path: str) -> Dict:
        """Extract all content from DOCX document"""
        try:
            doc = Document(file_path)
            
            content = {
                "paragraphs": [],
                "tables": [],
                "stats": {
                    "total_paragraphs": 0,
                    "total_tables": len(doc.tables),
                    "total_words": 0
                }
            }
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    para_data = {
                        "index": i,
                        "text": para.text.strip(),
                        "formatting": {
                            "bold": any(run.bold for run in para.runs if run.bold),
                            "italic": any(run.italic for run in para.runs if run.italic),
                            "alignment": str(para.alignment) if para.alignment else "LEFT"
                        },
                        "word_count": len(para.text.split())
                    }
                    content["paragraphs"].append(para_data)
                    content["stats"]["total_words"] += para_data["word_count"]
            
            content["stats"]["total_paragraphs"] = len(content["paragraphs"])
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = {
                    "index": i,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "cells": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_data.append({
                            "text": cell_text,
                            "row": row_idx,
                            "col": col_idx
                        })
                        content["stats"]["total_words"] += len(cell_text.split())
                    table_data["cells"].append(row_data)
                
                content["tables"].append(table_data)
            
            logger.info(f"Document extracted: {content['stats']['total_paragraphs']} paragraphs, "
                       f"{content['stats']['total_tables']} tables, "
                       f"{content['stats']['total_words']} words")
            
            return content
            
        except Exception as e:
            logger.error(f"Error extracting document: {e}")
            raise
    
    def translate_text(self, text: str, context: str = "") -> str:
        """Translate a piece of text using LLM"""
        if not text.strip():
            return text
            
        prompt = f"""Translate the following {self.source_language} text to {self.target_language}.

IMPORTANT RULES:
1. Maintain academic tone and precision
2. Preserve technical terminology accuracy  
3. Keep the same meaning and context
4. Return ONLY the translated text, no explanations
5. Do not add any additional text or notes

{context}

Text to translate:
{text}

Translation in {self.target_language}:"""
        
        try:
            response = completion(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=self.temperature,
                max_tokens=min(len(text.split()) * 3, 2000)  # Estimate tokens needed
            )
            
            translated = response.choices[0].message.content.strip()
            self.total_tokens += response.usage.total_tokens
            
            logger.debug(f"Translated {len(text.split())} words using {response.usage.total_tokens} tokens")
            return translated
            
        except Exception as e:
            logger.error(f"Translation error: {e}")
            logger.error(f"Failed text: {text[:100]}...")
            return text  # Return original if translation fails
    
    def translate_content(self, content: Dict) -> Dict:
        """Translate all extracted content"""
        translated_content = {
            "paragraphs": [],
            "tables": [],
            "stats": content["stats"].copy()
        }
        
        total_items = len(content["paragraphs"]) + len(content["tables"])
        current_item = 0
        
        # Translate paragraphs
        logger.info(f"Translating {len(content['paragraphs'])} paragraphs...")
        for para in content["paragraphs"]:
            current_item += 1
            logger.info(f"Progress: {current_item}/{total_items} - Translating paragraph {para['index']}")
            
            translated_para = para.copy()
            translated_para["text"] = self.translate_text(
                para["text"], 
                context="This is part of an academic document."
            )
            translated_content["paragraphs"].append(translated_para)
        
        # Translate tables
        if content["tables"]:
            logger.info(f"Translating {len(content['tables'])} tables...")
            for table in content["tables"]:
                current_item += 1
                logger.info(f"Progress: {current_item}/{total_items} - Translating table {table['index']}")
                
                translated_table = table.copy()
                translated_table["cells"] = []
                
                for row_idx, row in enumerate(table["cells"]):
                    translated_row = []
                    for cell in row:
                        translated_cell = cell.copy()
                        if cell["text"]:
                            translated_cell["text"] = self.translate_text(
                                cell["text"],
                                context="This is a table cell in an academic document."
                            )
                        translated_row.append(translated_cell)
                    translated_table["cells"].append(translated_row)
                
                translated_content["tables"].append(translated_table)
        
        translated_content["stats"]["total_tokens_used"] = self.total_tokens
        return translated_content
    
    def create_translated_document(self, translated_content: Dict, output_path: str):
        """Create new DOCX document with translated content"""
        try:
            doc = Document()
            
            # Add paragraphs in order
            for para in translated_content["paragraphs"]:
                p = doc.add_paragraph()
                run = p.add_run(para["text"])
                
                # Apply formatting
                if para["formatting"]["bold"]:
                    run.bold = True
                if para["formatting"]["italic"]:
                    run.italic = True
                
                # Set alignment
                alignment_map = {
                    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                
                alignment = para["formatting"]["alignment"]
                if alignment in alignment_map:
                    p.alignment = alignment_map[alignment]
            
            # Add tables
            for table_data in translated_content["tables"]:
                if table_data["cells"]:
                    rows = len(table_data["cells"])
                    cols = len(table_data["cells"][0]) if rows > 0 else 0
                    
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for row_idx, row in enumerate(table_data["cells"]):
                            for col_idx, cell_data in enumerate(row):
                                if row_idx < rows and col_idx < cols:
                                    table.cell(row_idx, col_idx).text = cell_data["text"]
            
            # Save document
            doc.save(output_path)
            logger.info(f"Translated document saved: {output_path}")
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            raise
    
    def translate_document(self, input_file: str, output_file: str = None) -> Dict:
        """Main translation function"""
        start_time = time.time()
        
        try:
            # Validate input
            input_path = Path(input_file)
            if not input_path.exists():
                raise FileNotFoundError(f"Input file not found: {input_file}")
            
            if not input_file.lower().endswith('.docx'):
                raise ValueError("Input file must be a DOCX document")
            
            # Set output path
            if not output_file:
                output_file = f"{input_path.stem}_translated_{self.target_language.lower()}.docx"
            
            logger.info(f"Starting translation: {input_file} -> {output_file}")
            logger.info(f"Language: {self.source_language} -> {self.target_language}")
            logger.info(f"Model: {self.model}")
            
            # Extract content
            logger.info("Step 1: Extracting document content...")
            content = self.extract_document_content(input_file)
            
            # Translate content
            logger.info("Step 2: Translating content...")
            translated_content = self.translate_content(content)
            
            # Create output document
            logger.info("Step 3: Creating translated document...")
            self.create_translated_document(translated_content, output_file)
            
            # Calculate stats
            execution_time = time.time() - start_time
            output_path = Path(output_file)
            
            result = {
                "status": "success",
                "input_file": str(input_path),
                "output_file": str(output_path),
                "source_language": self.source_language,
                "target_language": self.target_language,
                "execution_time_seconds": round(execution_time, 2),
                "file_size_bytes": output_path.stat().st_size if output_path.exists() else 0,
                "statistics": {
                    "paragraphs_translated": len(translated_content["paragraphs"]),
                    "tables_translated": len(translated_content["tables"]),
                    "total_words": translated_content["stats"]["total_words"],
                    "total_tokens_used": self.total_tokens
                }
            }
            
            logger.info("="*60)
            logger.info("TRANSLATION COMPLETED SUCCESSFULLY!")
            logger.info("="*60)
            logger.info(f"Input: {input_file}")
            logger.info(f"Output: {output_file}")
            logger.info(f"Language: {self.source_language} -> {self.target_language}")
            logger.info(f"Time: {execution_time:.2f} seconds")
            logger.info(f"Paragraphs: {result['statistics']['paragraphs_translated']}")
            logger.info(f"Tables: {result['statistics']['tables_translated']}")
            logger.info(f"Tokens used: {self.total_tokens}")
            logger.info("="*60)
            
            return result
            
        except Exception as e:
            execution_time = time.time() - start_time
            logger.error(f"Translation failed after {execution_time:.2f} seconds: {e}")
            
            return {
                "status": "error",
                "error": str(e),
                "execution_time_seconds": round(execution_time, 2),
                "input_file": input_file
            }


def main():
    """Main CLI function"""
    parser = argparse.ArgumentParser(
        description="CrowdWisdomTrading Direct Document Translator",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s academic_paper.docx                    # Translate to Spanish (default)
  %(prog)s paper.docx -t French                  # Translate to French
  %(prog)s doc.docx -t German -o german_doc.docx # Specify output file
  %(prog)s document.docx --verbose               # Verbose logging
        """
    )
    
    parser.add_argument("input_file", help="Input DOCX file to translate")
    parser.add_argument("-t", "--target", dest="target_language", default="Spanish", 
                       help="Target language (default: Spanish)")
    parser.add_argument("-o", "--output", dest="output_file", 
                       help="Output file path (default: auto-generated)")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging")
    parser.add_argument("--test", action="store_true", help="Test API connection")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        logger.debug("Verbose logging enabled")
    
    # Check API key
    api_key = os.getenv('GEMINI_API_KEY') or os.getenv('GOOGLE_API_KEY')
    if not api_key:
        print("Error: No API key found.")
        print("Set GEMINI_API_KEY or GOOGLE_API_KEY environment variable.")
        return 1
    
    if args.test:
        # Test API connection
        try:
            logger.info("Testing API connection...")
            response = completion(
                model="gemini/gemini-1.5-flash",
                messages=[{"role": "user", "content": "Translate 'Hello world' to Spanish"}],
                max_tokens=50
            )
            print(f"API test successful: {response.choices[0].message.content}")
            return 0
        except Exception as e:
            print(f"API test failed: {e}")
            return 1
    
    # Validate input file
    if not Path(args.input_file).exists():
        print(f"Error: Input file not found: {args.input_file}")
        return 1
    
    print("🚀 CrowdWisdomTrading Direct Translation System")
    print("="*60)
    
    # Create translator and run
    translator = DirectDocumentTranslator(target_language=args.target_language)
    result = translator.translate_document(args.input_file, args.output_file)
    
    if result['status'] == 'success':
        print(f"\n✅ Translation completed successfully!")
        print(f"📄 Input: {result['input_file']}")
        print(f"💾 Output: {result['output_file']}")
        print(f"🌐 Language: {result['source_language']} → {result['target_language']}")
        print(f"⏱️ Time: {result['execution_time_seconds']} seconds")
        print(f"📊 Stats:")
        print(f"   • Paragraphs: {result['statistics']['paragraphs_translated']}")
        print(f"   • Tables: {result['statistics']['tables_translated']}")
        print(f"   • Words: {result['statistics']['total_words']}")
        print(f"   • Tokens: {result['statistics']['total_tokens_used']}")
        print(f"📁 File size: {result['file_size_bytes']} bytes")
        return 0
    else:
        print(f"\n❌ Translation failed!")
        print(f"Error: {result['error']}")
        return 1


if __name__ == "__main__":
    sys.exit(main())