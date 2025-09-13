#!/usr/bin/env python3
"""
Resumable Document Translator - Can resume after quota limits
"""

import os
import json
import logging
import argparse
import sys
import time
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import litellm
from litellm import completion

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('resumable_translation.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class ResumableTranslator:
    def __init__(self, target_language="Spanish"):
        self.target_language = target_language
        self.source_language = "English"
        self.model = "gemini/gemini-1.5-flash"
        self.temperature = 0.3
        self.total_tokens = 0
        self.checkpoint_file = "translation_checkpoint.json"
        
    def save_checkpoint(self, data: Dict):
        """Save translation progress"""
        with open(self.checkpoint_file, 'w') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        logger.info(f"Checkpoint saved: {self.checkpoint_file}")
    
    def load_checkpoint(self) -> Dict:
        """Load previous translation progress"""
        if Path(self.checkpoint_file).exists():
            with open(self.checkpoint_file, 'r') as f:
                return json.load(f)
        return None
    
    def extract_document_content(self, file_path: str) -> Dict:
        """Extract all content from DOCX document"""
        try:
            doc = Document(file_path)
            
            content = {
                "paragraphs": [],
                "tables": [],
                "stats": {
                    "total_paragraphs": 0,
                    "total_tables": len(doc.tables),
                    "total_words": 0
                }
            }
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    para_data = {
                        "index": i,
                        "text": para.text.strip(),
                        "formatting": {
                            "bold": any(run.bold for run in para.runs if run.bold),
                            "italic": any(run.italic for run in para.runs if run.italic),
                            "alignment": str(para.alignment) if para.alignment else "LEFT"
                        },
                        "word_count": len(para.text.split())
                    }
                    content["paragraphs"].append(para_data)
                    content["stats"]["total_words"] += para_data["word_count"]
            
            content["stats"]["total_paragraphs"] = len(content["paragraphs"])
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = {
                    "index": i,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "cells": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_data.append({
                            "text": cell_text,
                            "row": row_idx,
                            "col": col_idx
                        })
                        content["stats"]["total_words"] += len(cell_text.split())
                    table_data["cells"].append(row_data)
                
                content["tables"].append(table_data)
            
            logger.info(f"Document extracted: {content['stats']['total_paragraphs']} paragraphs, "
                       f"{content['stats']['total_tables']} tables, "
                       f"{content['stats']['total_words']} words")
            
            return content
            
        except Exception as e:
            logger.error(f"Error extracting document: {e}")
            raise
    
    def translate_text_batch(self, texts: List[str], max_retries=3) -> List[str]:
        """Translate multiple texts in one request to save quota"""
        if not texts:
            return texts
            
        # Combine texts for batch translation
        combined_text = "\n\n---SEPARATOR---\n\n".join(texts)
        
        prompt = f"""Translate the following {self.source_language} texts to {self.target_language}.

IMPORTANT RULES:
1. Maintain academic tone and precision
2. Preserve technical terminology accuracy  
3. Keep the same meaning and context
4. Return ONLY the translated texts separated by ---SEPARATOR---
5. Do not add any additional text or notes
6. Maintain the exact same number of text sections

Texts to translate (separated by ---SEPARATOR---):
{combined_text}

Translation in {self.target_language} (maintain separators):"""
        
        for attempt in range(max_retries):
            try:
                response = completion(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=self.temperature,
                    max_tokens=min(len(combined_text.split()) * 2, 4000)
                )
                
                translated = response.choices[0].message.content.strip()
                self.total_tokens += response.usage.total_tokens
                
                # Split back into individual translations
                translations = translated.split("---SEPARATOR---")
                translations = [t.strip() for t in translations]
                
                # If we got the right number of translations, return them
                if len(translations) == len(texts):
                    logger.info(f"Batch translated {len(texts)} items using {response.usage.total_tokens} tokens")
                    return translations
                else:
                    logger.warning(f"Translation count mismatch: {len(translations)} vs {len(texts)}")
                    # Fall back to individual translation
                    break
                    
            except Exception as e:
                if "quota" in str(e).lower() or "rate limit" in str(e).lower():
                    logger.error(f"Quota exceeded: {e}")
                    raise e  # Don't retry quota errors
                elif attempt < max_retries - 1:
                    logger.warning(f"Translation attempt {attempt + 1} failed: {e}. Retrying...")
                    time.sleep(2)  # Wait before retry
                else:
                    logger.error(f"Translation failed after {max_retries} attempts: {e}")
                    return texts  # Return originals if all attempts fail
        
        # Fallback: translate individually
        return [self.translate_single_text(text) for text in texts]
    
    def translate_single_text(self, text: str) -> str:
        """Translate a single piece of text"""
        if not text.strip():
            return text
            
        prompt = f"""Translate this {self.source_language} text to {self.target_language}:

{text}

Translation:"""
        
        try:
            response = completion(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=self.temperature,
                max_tokens=min(len(text.split()) * 2, 2000)
            )
            
            translated = response.choices[0].message.content.strip()
            self.total_tokens += response.usage.total_tokens
            return translated
            
        except Exception as e:
            logger.error(f"Single translation error: {e}")
            return text  # Return original if translation fails
    
    def translate_document(self, input_file: str, output_file: str = None, batch_size: int = 5) -> Dict:
        """Main translation function with resume capability"""
        start_time = time.time()
        
        try:
            input_path = Path(input_file)
            if not input_path.exists():
                raise FileNotFoundError(f"Input file not found: {input_file}")
            
            if not output_file:
                output_file = f"{input_path.stem}_translated_{self.target_language.lower()}.docx"
            
            logger.info(f"Starting resumable translation: {input_file} -> {output_file}")
            logger.info(f"Language: {self.source_language} -> {self.target_language}")
            logger.info(f"Batch size: {batch_size}")
            
            # Try to load checkpoint
            checkpoint = self.load_checkpoint()
            
            if checkpoint and checkpoint.get('input_file') == str(input_path):
                logger.info(f"Resuming from checkpoint: {checkpoint['completed_items']} items completed")
                content = checkpoint['content']
                translated_content = checkpoint['translated_content']
                start_index = checkpoint['completed_items']
            else:
                logger.info("Starting fresh translation...")
                content = self.extract_document_content(input_file)
                translated_content = {
                    "paragraphs": [],
                    "tables": [],
                    "stats": content["stats"].copy()
                }
                start_index = 0
            
            total_paragraphs = len(content["paragraphs"])
            
            # Process paragraphs in batches
            logger.info(f"Processing {total_paragraphs - start_index} remaining paragraphs...")
            
            for i in range(start_index, total_paragraphs, batch_size):
                batch_end = min(i + batch_size, total_paragraphs)
                batch = content["paragraphs"][i:batch_end]
                
                logger.info(f"Translating batch {i//batch_size + 1}: paragraphs {i}-{batch_end-1}")
                
                try:
                    # Extract texts for batch translation
                    texts = [para["text"] for para in batch]
                    translated_texts = self.translate_text_batch(texts)
                    
                    # Create translated paragraph objects
                    for j, (para, translated_text) in enumerate(zip(batch, translated_texts)):
                        translated_para = para.copy()
                        translated_para["text"] = translated_text
                        translated_content["paragraphs"].append(translated_para)
                    
                    # Save checkpoint after each batch
                    checkpoint_data = {
                        "input_file": str(input_path),
                        "target_language": self.target_language,
                        "completed_items": batch_end,
                        "total_items": total_paragraphs,
                        "content": content,
                        "translated_content": translated_content,
                        "total_tokens_used": self.total_tokens,
                        "timestamp": time.time()
                    }
                    self.save_checkpoint(checkpoint_data)
                    
                    logger.info(f"Progress: {batch_end}/{total_paragraphs} paragraphs completed")
                    
                except Exception as e:
                    if "quota" in str(e).lower() or "rate limit" in str(e).lower():
                        logger.error(f"Quota exceeded at paragraph {i}. Translation paused.")
                        logger.info(f"Resume tomorrow with: python {sys.argv[0]} --resume")
                        return {
                            "status": "paused",
                            "reason": "quota_exceeded",
                            "completed_items": i,
                            "total_items": total_paragraphs,
                            "checkpoint_file": self.checkpoint_file
                        }
                    else:
                        raise e
            
            # Process tables if any
            if content["tables"]:
                logger.info(f"Processing {len(content['tables'])} tables...")
                for table in content["tables"]:
                    # For now, just copy tables (can be enhanced later)
                    translated_content["tables"].append(table)
            
            # Create final document
            logger.info("Creating final document...")
            self.create_translated_document(translated_content, output_file)
            
            # Clean up checkpoint
            if Path(self.checkpoint_file).exists():
                Path(self.checkpoint_file).unlink()
                logger.info("Checkpoint file cleaned up")
            
            execution_time = time.time() - start_time
            output_path = Path(output_file)
            
            result = {
                "status": "success",
                "input_file": str(input_path),
                "output_file": str(output_path),
                "source_language": self.source_language,
                "target_language": self.target_language,
                "execution_time_seconds": round(execution_time, 2),
                "file_size_bytes": output_path.stat().st_size if output_path.exists() else 0,
                "statistics": {
                    "paragraphs_translated": len(translated_content["paragraphs"]),
                    "tables_translated": len(translated_content["tables"]),
                    "total_words": translated_content["stats"]["total_words"],
                    "total_tokens_used": self.total_tokens
                }
            }
            
            logger.info("="*60)
            logger.info("TRANSLATION COMPLETED SUCCESSFULLY!")
            logger.info("="*60)
            logger.info(f"Output: {output_file}")
            logger.info(f"Time: {execution_time:.2f} seconds")
            logger.info(f"Tokens used: {self.total_tokens}")
            logger.info("="*60)
            
            return result
            
        except Exception as e:
            execution_time = time.time() - start_time
            logger.error(f"Translation failed: {e}")
            
            return {
                "status": "error",
                "error": str(e),
                "execution_time_seconds": round(execution_time, 2)
            }
    
    def create_translated_document(self, translated_content: Dict, output_path: str):
        """Create new DOCX document with translated content"""
        try:
            doc = Document()
            
            # Add paragraphs
            for para in translated_content["paragraphs"]:
                p = doc.add_paragraph()
                run = p.add_run(para["text"])
                
                # Apply formatting
                if para["formatting"]["bold"]:
                    run.bold = True
                if para["formatting"]["italic"]:
                    run.italic = True
                
                # Set alignment
                alignment_map = {
                    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                
                alignment = para["formatting"]["alignment"]
                if alignment in alignment_map:
                    p.alignment = alignment_map[alignment]
            
            # Add tables
            for table_data in translated_content["tables"]:
                if table_data["cells"]:
                    rows = len(table_data["cells"])
                    cols = len(table_data["cells"][0]) if rows > 0 else 0
                    
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for row_idx, row in enumerate(table_data["cells"]):
                            for col_idx, cell_data in enumerate(row):
                                if row_idx < rows and col_idx < cols:
                                    table.cell(row_idx, col_idx).text = cell_data["text"]
            
            doc.save(output_path)
            logger.info(f"Document saved: {output_path}")
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            raise


def main():
    parser = argparse.ArgumentParser(description="Resumable Document Translator")
    
    parser.add_argument("input_file", nargs='?', help="Input DOCX file")
    parser.add_argument("-t", "--target", dest="target_language", default="Spanish", help="Target language")
    parser.add_argument("-o", "--output", dest="output_file", help="Output file")
    parser.add_argument("--resume", action="store_true", help="Resume from checkpoint")
    parser.add_argument("--batch-size", type=int, default=3, help="Batch size (default: 3)")
    parser.add_argument("--status", action="store_true", help="Show checkpoint status")
    
    args = parser.parse_args()
    
    translator = ResumableTranslator(target_language=args.target_language)
    
    if args.status:
        checkpoint = translator.load_checkpoint()
        if checkpoint:
            print(f"Checkpoint found:")
            print(f"  File: {checkpoint['input_file']}")
            print(f"  Progress: {checkpoint['completed_items']}/{checkpoint['total_items']}")
            print(f"  Language: {checkpoint['target_language']}")
            print(f"  Tokens used: {checkpoint['total_tokens_used']}")
            print(f"  Date: {time.ctime(checkpoint['timestamp'])}")
        else:
            print("No checkpoint found")
        return
    
    if args.resume:
        checkpoint = translator.load_checkpoint()
        if checkpoint:
            print("Resuming translation from checkpoint...")
            result = translator.translate_document(
                checkpoint['input_file'], 
                None,  # Use default output name
                args.batch_size
            )
        else:
            print("No checkpoint found to resume from")
            return
    else:
        if not args.input_file:
            print("Usage: python resumable_translator.py input_file.docx -t Spanish")
            print("       python resumable_translator.py --resume")
            print("       python resumable_translator.py --status")
            return
        
        result = translator.translate_document(args.input_file, args.output_file, args.batch_size)
    
    if result['status'] == 'success':
        print("Translation completed successfully!")
        print(f"Output: {result['output_file']}")
    elif result['status'] == 'paused':
        print(f"Translation paused due to quota limit.")
        print(f"Completed: {result['completed_items']}/{result['total_items']} items")
        print("Resume tomorrow with: python resumable_translator.py --resume")
    else:
        print(f"Translation failed: {result['error']}")


if __name__ == "__main__":
    main()